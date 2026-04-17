from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document as Docx
import time
import os

import streamlit as st

from agent.react_agent import ReactAgent
from rag.kdb_service import KDBService
from langchain_core.documents import Document
from utils.logger_handler import logger
from utils.file_handler import get_file_md5_hex

# 初始化数据库服务，确保在会话状态中只创建一次
if "db_service" not in st.session_state:
    st.session_state["db_service"] = KDBService()

# 页面配置，设置标题、图标和布局
st.set_page_config(
    page_title="智能客服系统",
    page_icon="🤖",
    layout="wide"
)

# 创建侧边栏导航，让用户可以在不同功能页面间切换
page = st.sidebar.radio(
    "导航",
    ["对话", "知识库管理"]
)

# 对话页面功能
if page == "对话":
    # 标题
    st.title("智能客服")
    st.divider()

    # 初始化智能体，确保在会话状态中只创建一次
    if "agent" not in st.session_state:
        st.session_state["agent"] = ReactAgent()

    # 初始化消息列表，用于存储对话历史，注意这里仅仅是短期对话历史。
    if "message" not in st.session_state:
        st.session_state["message"] = []

    # 显示历史消息
    for msg in st.session_state["message"]:
        st.chat_message(msg["role"]).write(msg["content"])

    # 用户输入提示词
    prompt = st.chat_input("请输入你的问题")

    # 处理用户输入
    if prompt:
        # 显示用户消息
        st.chat_message("user").write(prompt)
        # 将用户消息添加到历史记录
        st.session_state["message"].append({"role": "user", "content": prompt})

        # 用于存储响应消息
        response_messages = []
        # 显示加载动画
        with st.spinner("思考中..."): # 到目前为止都是简单的前端
            # 调用智能体执行流式响应
            res_stream = st.session_state["agent"].execute_stream(prompt) # 前端开始调用后端。

            # 定义捕获函数，用于处理生成器输出并模拟打字效果
            def capture(generator, cache_list):
                for chunk in generator:
                    cache_list.append(chunk)

                    for char in chunk: 
                        time.sleep(0.01) # 模拟延迟
                        yield char

            # 显示助手的响应，使用流式输出
            st.chat_message("assistant").write_stream(capture(res_stream, response_messages))
            # 将助手响应添加到历史记录
            st.session_state["message"].append({"role": "assistant", "content": response_messages[-1]})
            # 重新运行页面以更新显示
            st.rerun()  # 重新运行

# 知识库管理页面功能
elif page == "知识库管理":
    # 页面标题
    st.title("知识库管理")
    st.divider()
    
    # 创建两个选项卡，分别对应不同的知识库操作
    #tab1, tab2 = st.tabs(["添加文档", "已有文档操作"])

    # 修改后
    tab1, tab2, tab3 = st.tabs(["添加文档", "已有文档操作", "调试专用"])
    
    # 添加文档功能
    with tab1:
        # 子标题
        st.subheader("添加新文档")
        
        # 文档内容输入区域
        content = st.text_area("文档内容", height=200, key="add_content")

        # 高级选项区域，用于输入文档元数据
        with st.expander("高级选项（元数据）"):
            # 创建元数据字典
            metadata = {}
            # 输入文档标题
            metadata["title"] = st.text_input("标题", key="add_title")
            # 输入文档来源
            metadata["source"] = st.text_input("来源", key="add_source")
            # 输入文档作者
            metadata["author"] = st.text_input("作者", key="add_author")
            # 输入文档分类
            metadata["category"] = st.text_input("分类")
        
        # 添加文档按钮
        if st.button("添加文档", type="primary"):
            # 检查是否输入了文档内容
            if content:
                # 生成文档ID
                import uuid
                main_id = str(uuid.uuid4())
                metadata["main_id"] = main_id
                
                # 记录日志
                logger.info(f"[前端]准备添加文档，main_id: {main_id}")
                
                # 调用数据库服务添加文档
                doc = Document(page_content=content, metadata=metadata)
                doc_ids = st.session_state["db_service"].add_documents([doc])
                
                # 检查添加是否成功
                if doc_ids:
                    # 显示成功消息
                    st.success(f"文档添加成功！文档ID: {main_id}")
                    logger.info(f"[前端]文档添加成功，main_id: {main_id}, 生成碎片数: {len(doc_ids)}")
                else:
                    # 显示失败消息
                    st.error("文档添加失败！")
                    logger.error(f"[前端]文档添加失败，main_id: {main_id}")
            else:
                # 提示用户输入内容
                st.warning("请输入文档内容")
        
        # 分隔线
        st.divider()
        # 批量添加文档子标题
        st.subheader("批量添加文档")
        
        # 文件上传组件，支持多文件上传
        uploaded_files = st.file_uploader(
            "上传文件",
            type=["txt", "pdf", "docx"],
            accept_multiple_files=True,
            help="支持上传txt、pdf、docx格式的文件"
        )
        
        def extract_text_from_file(file):
            """从上传的文件中提取文本内容"""
            try:
                if file.type == "application/pdf":
                    # 处理PDF文件
                    pdf_reader = PdfReader(BytesIO(file.read()))
                    content = ""
                    for page in pdf_reader.pages:
                        content += page.extract_text() + "\n"
                    return content
                elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    # 处理DOCX文件
                    doc = Docx(BytesIO(file.read()))
                    content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    return content
                else:
                    # 处理TXT文件
                    return file.read().decode("utf-8")
            except Exception as e:
                logger.error(f"[前端]解析文件 {file.name} 时出错: {str(e)}", exc_info=True)
                raise


        # 批量添加按钮
        if st.button("批量添加", type="primary"):
            # 检查是否上传了文件
            if uploaded_files:
                # 初始化成功计数器
                success_count = 0
                # 遍历上传的文件
                for file in uploaded_files:
                    try:
                        # 读取文件内容
                        content = extract_text_from_file(file)
                        
                        # 将文件名拼接到内容开头，以便支持通过文件名搜索
                        # 格式：文件名: xxx.txt\n内容:\n...
                        content = f"文件名: {file.name}\n内容:\n{content}"
                        
                        # 生成文档ID
                        import uuid
                        main_id = str(uuid.uuid4())
                        
                        # 创建文件元数据
                        metadata = {
                            "title": file.name,
                            "source": file.name,
                            "file_type": file.type,
                            "main_id": main_id
                        }
                        
                        # 记录日志
                        logger.info(f"[前端]准备批量添加文件，文件名: {file.name}, main_id: {main_id}")
                        
                        # 调用数据库服务添加文档
                        doc = Document(page_content=content, metadata=metadata)
                        doc_ids = st.session_state["db_service"].add_documents([doc])
                        
                        # 检查添加是否成功
                        if doc_ids:
                            success_count += 1
                            # 记录成功日志
                            logger.info(f"[前端]文件 {file.name} 添加成功，main_id: {main_id}, 生成碎片数: {len(doc_ids)}")
                        else:
                            # 记录失败日志
                            logger.error(f"[前端]文件 {file.name} 添加失败，main_id: {main_id}")
                    except Exception as e:
                        # 记录异常日志
                        logger.error(f"[前端]处理文件 {file.name} 时出错: {str(e)}", exc_info=True)
                
                # 显示批量添加结果
                st.success(f"已添加 {success_count}/{len(uploaded_files)} 个文件")
                logger.info(f"[前端]批量添加完成，成功: {success_count}/{len(uploaded_files)}")
            else:
                # 提示用户上传文件
                st.warning("请先上传文件")
    
    # 已有文档操作功能
    with tab2:

        # 子标题
        st.subheader("搜索文档")

        # 搜索关键词输入
        query = st.text_input("搜索关键词（文档内容或标题）", key="search_query")
        # 返回结果数量滑块
        k = st.slider("返回结果数量", min_value=1, max_value=10, value=3)

        # 搜索按钮
        if st.button("搜索", type="primary"):
            # 检查是否输入了搜索关键词
            if query:
                # 记录搜索日志
                logger.info(f"[前端]开始搜索文档，关键词: {query}")

                # 调用数据库服务搜索文档
                results = st.session_state["db_service"].search_documents(query, k)

                # 将搜索结果存储到会话状态中。因为streamlit的按钮，只要按一次，就会自动触发重运行。
                # 这种重新运行会将临时的搜索结果清空，所有我们必须在重运行之前，将结果存放到会话状态里，
                #   这样就可以在重运行之后，从会话状态中取出结果显示在界面上。
                st.session_state["search_results"] = results

            else:
                # 提示输入搜索关键词
                st.warning("请输入搜索关键词")

        # ==============================================
        # 【核心】从session_state读取结果，独立于搜索按钮之外
        # ==============================================
        search_results = st.session_state.get("search_results", [])

        # 显示搜索结果（每次重运行都渲染，不会消失）
        if search_results:
            st.success(f"找到 {len(search_results)} 条相关文档")
            logger.info(f"[前端]搜索完成，找到 {len(search_results)} 条相关文档")

            # 遍历搜索结果
            for i, doc in enumerate(search_results, 1):
                # 获取文档ID
                main_id = doc.metadata.get('main_id', '')
                # 创建列布局
                col1, col2, col3 = st.columns([3, 1, 1])

                logger.info(f"[前端]显示搜索结果，main_id: {main_id}")
                # 第一列显示文档信息
                with col1:
                    # 使用可折叠区域显示每个文档
                    with st.expander(f"文档 {i} - {doc.metadata.get('title', '未知标题')}"):
                        # 显示文档内容
                        st.text_area("内容", doc.page_content, height=150, key=f"search_{i}")

                        # 显示文档元数据
                        if doc.metadata:
                            st.json(doc.metadata)

                # 第二列显示查看按钮
                with col2:
                    if st.button("查看", key=f"view_{i}"):
                        # 记录查看日志
                        logger.info(f"[前端]用户请求查看文档，main_id: {main_id}")

                        # 获取完整文档
                        full_doc = st.session_state["db_service"].get_document_by_main_id(main_id)

                        if full_doc:
                            # 显示完整文档
                            st.info(f"完整文档 (main_id: {main_id})")
                            st.text_area("完整内容", full_doc.page_content, height=300, key=f"full_{i}")
                            st.json(full_doc.metadata)
                            logger.info(f"[前端]成功获取完整文档，main_id: {main_id}")
                        else:
                            st.error("获取完整文档失败")
                            logger.error(f"[前端]获取完整文档失败，main_id: {main_id}")

                # 第三列显示删除和更新按钮
                with col3:
                    # 删除按钮
                    if st.button("删除", key=f"delete_{i}"):
                        logger.info(f"[前端]用户请求删除文档，main_id: {main_id}")
                        all_docs = st.session_state["db_service"].get_all_documents()
                        slice_ids = [d.metadata.get('slice_id') for d in all_docs if d.metadata.get('main_id') == main_id]
                        if slice_ids and st.session_state["db_service"].delete_documents(slice_ids):
                            st.success(f"文档 {main_id} 及其所有碎片已删除")
                            st.session_state["search_results"] = []
                            time.sleep(2.5)  # 停留2.5秒再刷新，以防用户看不到删除成功的提示
                            st.rerun()
                        else:
                            st.error(f"删除文档 {main_id} 失败")

                    # 更新按钮（平级，不嵌套）
                    if st.button("更新", key=f"update_{i}"):
                        # 把编辑状态存入会话
                        st.session_state["editing_main_id"] = main_id
                        st.session_state["editing_index"] = i
                        st.rerun()

                # ===================== 【关键】把更新表单移到按钮外面 =====================
                # 独立渲染，动用持久化后的编辑状态信息，不嵌套在任何按钮里
                if "editing_main_id" in st.session_state and st.session_state["editing_main_id"] == main_id:
                    i = st.session_state["editing_index"]
                    full_doc = st.session_state["db_service"].get_document_by_main_id(main_id)
                    
                    if full_doc:
                        st.info(f"更新文档 (main_id: {main_id})")
                        updated_content = st.text_area("文档内容", full_doc.page_content, height=200, key=f"update_content_{i}")
                        
                        with st.expander("元数据"):
                            updated_metadata = full_doc.metadata.copy()
                            updated_metadata["title"] = st.text_input("标题", updated_metadata.get("title", ""), key=f"update_title_{i}")
                            updated_metadata["source"] = st.text_input("来源", updated_metadata.get("source", ""), key=f"update_source_{i}")
                            updated_metadata["author"] = st.text_input("作者", updated_metadata.get("author", ""), key=f"update_author_{i}")
                            updated_metadata["category"] = st.text_input("分类", updated_metadata.get("category", ""), key=f"update_category_{i}")
                        
                        # 确认更新按钮（现在是平级，能正常触发）
                        if st.button("确认更新", key=f"confirm_update_{i}"):
                            logger.info(f"[前端]用户确认更新文档，main_id: {main_id}")
                            
                            all_docs = st.session_state["db_service"].get_all_documents()
                            slice_ids = [d.metadata.get('slice_id') for d in all_docs if d.metadata.get('main_id') == main_id]
                            
                            if slice_ids and st.session_state["db_service"].delete_documents(slice_ids):
                                logger.info(f"[前端]已删除旧文档碎片")
                                updated_doc = Document(page_content=updated_content, metadata=updated_metadata)
                                new_doc_ids = st.session_state["db_service"].add_documents([updated_doc])
                                
                                logger.info(f"[前端]已添加新文档碎片, 最好确认一下#########")
                                if new_doc_ids:
                                    st.success(f"文档 {main_id} 更新成功")
                                    logger.info(f"[前端]文档 {main_id} 更新成功")
                                    # 清空编辑状态
                                    del st.session_state["editing_main_id"]
                                    del st.session_state["editing_index"]
                                    st.session_state["search_results"] = []
                                    time.sleep(2.5)  # 停留2.5秒再刷新，以防用户看不到更新成功的提示
                                    st.rerun()
                                else:
                                    st.error("添加更新后的文档失败")
                                    logger.error(f"[前端]添加更新后的文档失败")
                            else:
                                st.error("删除旧文档碎片失败")

        # 无结果提示（独立渲染）
        elif "search_results" in st.session_state and not search_results:
            st.warning("未找到相关文档")
            logger.info(f"[前端]搜索完成，未找到相关文档，关键词: {query}")


    # 在已有文档操作功能结束后添加新的tab3内容
    with tab3:
        st.subheader("调试专用")
        st.warning("⚠️ 以下操作仅用于调试，请谨慎使用！")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("列出全部文档", type="secondary"):
                try:
                    all_docs = st.session_state["db_service"].get_all_documents()
                    if all_docs:
                        st.success(f"知识库中共有 {len(all_docs)} 个文档碎片")
                        for i, doc in enumerate(all_docs, 1):
                            with st.expander(f"文档碎片 {i} - slice_id: {doc.metadata.get('slice_id', 'N/A')}"):
                                st.text_area("内容", doc.page_content, height=100, key=f"debug_content_{i}")
                                st.json(doc.metadata)
                    else:
                        st.info("知识库中没有文档")
                    logger.info(f"[前端]调试：列出全部文档，共 {len(all_docs)} 个")
                except Exception as e:
                    st.error(f"列出文档失败: {str(e)}")
                    logger.error(f"[前端]调试：列出文档失败: {str(e)}", exc_info=True)
        
        with col2:
            if st.button("删除全部文档", type="primary"):
                try:
                    if st.session_state["db_service"].delete_all_documents():
                        st.success("已删除知识库中的所有文档")
                        logger.info("[前端]调试：已删除全部文档")
                        st.rerun()
                    else:
                        st.error("删除文档失败")
                        logger.error("[前端]调试：删除全部文档失败")
                except Exception as e:
                    st.error(f"删除文档失败: {str(e)}")
                    logger.error(f"[前端]调试：删除文档失败: {str(e)}", exc_info=True)


